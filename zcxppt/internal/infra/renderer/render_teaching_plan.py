#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word Teaching Plan Renderer

Receives JSON via stdin:
{
    "output_path": str,   # Output .docx path
    "plan": str          # JSON string of LLM-generated teaching plan
}

The plan JSON structure:
{
    "title": str,
    "subject": str,
    "grade": str,
    "duration": str,
    "teaching_goals": [...],
    "teaching_focus": [...],
    "teaching_difficulties": [...],
    "teaching_methods": [...],
    "teaching_aids": [...],
    "teaching_process": {
        "warm_up": {"duration": str, "content": str},
        "introduction": {"duration": str, "content": str},
        "new_teaching": [{"step": int, "title": str, "duration": str, "content": str, "activities": [...]}],
        "practice": {"duration": str, "content": str},
        "summary": {"duration": str, "content": str},
        "homework": [...]
    },
    "classroom_activities": [{"name": str, "type": str, "duration": str, "description": str, "purpose": str}],
    "teaching_reflection": str
}

Outputs JSON to stdout:
{
    "success": bool,
    "docx_path": str,
    "error": str
}
"""

import json
import sys
import traceback


def run():
    raw = sys.stdin.read()
    if not raw.strip():
        print(json.dumps({"success": False, "docx_path": "", "error": "empty input"}))
        sys.exit(1)

    try:
        req = json.loads(raw)
    except json.JSONDecodeError as e:
        print(json.dumps({"success": False, "docx_path": "", "error": f"invalid json: {e}"}))
        sys.exit(1)

    output_path = req.get("output_path", "teaching_plan.docx")
    plan_str = req.get("plan", "")

    if not plan_str:
        print(json.dumps({"success": False, "docx_path": "", "error": "plan field is required"}))
        sys.exit(1)

    try:
        plan = json.loads(plan_str)
    except json.JSONDecodeError as e:
        print(json.dumps({"success": False, "docx_path": "", "error": f"plan is not valid json: {e}"}))
        sys.exit(1)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print(json.dumps({"success": False, "docx_path": "", "error": "python-docx not installed: pip install python-docx"}))
        sys.exit(1)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Helper functions ──────────────────────────────────────────

    def add_heading(doc, text, level=1):
        """Add a styled heading."""
        p = doc.add_heading("", level=0)
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 78, 121)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(68, 114, 196)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_para(doc, text="", bold=False, indent=False, font_size=11, color=None):
        """Add a paragraph."""
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(1)
        if text:
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        return p

    def add_bullet(doc, text, level=0):
        """Add a bullet point."""
        style = "List Bullet" if level == 0 else "List Bullet 2"
        try:
            p = doc.add_paragraph(style=style)
        except Exception:
            p = doc.add_paragraph(style="List Bullet")
        if level > 0:
            p.paragraph_format.left_indent = Cm(1 + level * 0.5)
        run = p.add_run(text)
        run.font.size = Pt(11)
        return p

    def section_label(doc, label, content=None):
        """Add a labeled paragraph."""
        p = doc.add_paragraph()
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = RGBColor(31, 78, 121)
        if content:
            run_content = p.add_run(content)
            run_content.font.size = Pt(11)
        return p

    # ── Cover Page ────────────────────────────────────────────────
    add_heading(doc, plan.get("title", "教案"), level=1)
    add_para(doc)

    # Metadata row
    meta_parts = []
    if plan.get("subject"):
        meta_parts.append(f"学科：{plan['subject']}")
    if plan.get("grade"):
        meta_parts.append(f"年级：{plan['grade']}")
    if plan.get("duration"):
        meta_parts.append(f"课时：{plan['duration']}")
    if meta_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("    ".join(meta_parts))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(80, 80, 80)
    add_para(doc)

    # ── 一、教学目标 ──────────────────────────────────────────────
    add_heading(doc, "一、教学目标", level=2)
    goals = plan.get("teaching_goals", [])
    if goals:
        for goal in goals:
            add_bullet(doc, goal)
    else:
        add_para(doc, "（未提供）")

    # ── 二、教学重点与难点 ────────────────────────────────────────
    add_heading(doc, "二、教学重点与难点", level=2)

    focus = plan.get("teaching_focus", [])
    if focus:
        section_label(doc, "教学重点：")
        for f in focus:
            add_bullet(doc, f, level=1)

    difficulties = plan.get("teaching_difficulties", [])
    if difficulties:
        section_label(doc, "教学难点：")
        for d in difficulties:
            add_bullet(doc, d, level=1)

    # ── 三、教学方法 ──────────────────────────────────────────────
    add_heading(doc, "三、教学方法", level=2)
    methods = plan.get("teaching_methods", [])
    if methods:
        for m in methods:
            add_bullet(doc, m)
    else:
        add_para(doc, "（未提供）")

    # ── 四、教学准备 ──────────────────────────────────────────────
    aids = plan.get("teaching_aids", [])
    if aids:
        add_heading(doc, "四、教学准备", level=2)
        for aid in aids:
            add_bullet(doc, aid)

    # ── 五、教学过程 ──────────────────────────────────────────────
    add_heading(doc, "五、教学过程", level=2)
    process = plan.get("teaching_process", {})

    # 5.1 热身导入
    warm_up = process.get("warm_up", {})
    if warm_up.get("content"):
        add_heading(doc, "（一）热身导入", level=2)
        if warm_up.get("duration"):
            section_label(doc, "时长：", warm_up["duration"])
        add_para(doc, warm_up["content"])

    # 5.2 新课导入
    intro = process.get("introduction", {})
    if intro.get("content"):
        add_heading(doc, "（二）新课导入", level=2)
        if intro.get("duration"):
            section_label(doc, "时长：", intro["duration"])
        add_para(doc, intro["content"])

    # 5.3 新授环节
    new_teach = process.get("new_teaching", [])
    if new_teach:
        add_heading(doc, "（三）新授环节", level=2)
        for step_data in new_teach:
            # Step header
            step_num = step_data.get("step", "?")
            step_title = step_data.get("title", "")
            p = doc.add_paragraph()
            run = p.add_run(f"步骤 {step_num}：{step_title}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(31, 78, 121)

            if step_data.get("duration"):
                section_label(doc, "时长：", step_data["duration"])

            content = step_data.get("content", "")
            if content:
                add_para(doc, content, indent=True)

            for act in step_data.get("activities", []):
                add_bullet(doc, f"【活动】{act}", level=1)
            add_para(doc)

    # 5.4 课堂练习
    practice = process.get("practice", {})
    if practice.get("content"):
        add_heading(doc, "（四）课堂练习", level=2)
        if practice.get("duration"):
            section_label(doc, "时长：", practice["duration"])
        add_para(doc, practice["content"])

    # 5.5 课堂总结
    summary = process.get("summary", {})
    if summary.get("content"):
        add_heading(doc, "（五）课堂总结", level=2)
        if summary.get("duration"):
            section_label(doc, "时长：", summary["duration"])
        add_para(doc, summary["content"])

    # 5.6 课后作业
    homework = process.get("homework", [])
    if homework:
        add_heading(doc, "（六）课后作业", level=2)
        for i, hw in enumerate(homework, 1):
            add_para(doc, f"{i}. {hw}")

    # ── 六、课堂活动设计 ──────────────────────────────────────────
    activities = plan.get("classroom_activities", [])
    if activities:
        add_heading(doc, "六、课堂活动设计", level=2)
        for idx, act in enumerate(activities, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"活动 {idx}：{act.get('name', '')}（{act.get('type', '')}）")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(31, 78, 121)

            if act.get("duration"):
                section_label(doc, "时长：", act["duration"], indent=True)
            if act.get("description"):
                section_label(doc, "描述：", act["description"], indent=True)
            if act.get("purpose"):
                section_label(doc, "目的：", act["purpose"], indent=True)
            add_para(doc)

    # ── 七、教学反思 ──────────────────────────────────────────────
    reflection = plan.get("teaching_reflection", "")
    if reflection:
        add_heading(doc, "七、教学反思", level=2)
        add_para(doc, reflection)

    doc.save(output_path)
    print(json.dumps({"success": True, "docx_path": output_path, "error": ""}))


if __name__ == "__main__":
    try:
        run()
    except Exception as e:
        tb = traceback.format_exc()
        print(json.dumps({
            "success": False,
            "docx_path": "",
            "error": f"{type(e).__name__}: {e}\n{tb}"
        }))
        sys.exit(1)

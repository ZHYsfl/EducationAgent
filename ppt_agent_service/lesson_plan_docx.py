"""
根据任务与课件页生成教案 Word（.docx），供 POST /api/v1/ppt/export format=docx。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from ppt_agent_service.slide_py_code import extract_slide_html_from_py
from ppt_agent_service.task_manager import TaskState


def _html_to_plain_text(html: str, max_len: int = 4000) -> str:
    if not html:
        return ""
    t = re.sub(r"(?is)<script.*?>.*?</script>", " ", html)
    t = re.sub(r"(?is)<style.*?>.*?</style>", " ", t)
    t = re.sub(r"<br\s*/?>", "\n", t, flags=re.I)
    t = re.sub(r"</(p|div|h[1-6]|li|tr)>", "\n", t, flags=re.I)
    t = re.sub(r"<[^>]+>", " ", t)
    t = re.sub(r"[ \t\r\f\v]+", " ", t)
    t = "\n".join(line.strip() for line in t.split("\n"))
    t = re.sub(r"\n{3,}", "\n\n", t).strip()
    if len(t) > max_len:
        t = t[: max_len - 1] + "…"
    return t


def _te_lines(label: str, value: Any) -> list[tuple[str, str]]:
    if value is None:
        return []
    if isinstance(value, list):
        if not value:
            return []
        text = "；".join(str(x).strip() for x in value if str(x).strip())
        return [(label, text)] if text else []
    s = str(value).strip()
    return [(label, s)] if s else []


def build_lesson_plan_docx(task: TaskState, output_path: Path) -> None:
    """
    写入标准 .docx：教学要素 + 任务描述摘要 + 各页从幻灯片 HTML 抽取的文本要点。
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError("python-docx 未安装，无法导出教案") from e

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("教案", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading((task.topic or "（未命名课程）").strip(), level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("基本信息")
    r.bold = True
    r.font.size = Pt(12)

    meta_lines = [
        ("任务 ID", task.task_id),
        ("会话 ID", task.session_id),
        ("授课对象", task.audience or "—"),
        ("全局风格", task.global_style or "—"),
        ("课件页数", str(len(task.page_order))),
        ("期望页数（初始化）", str(task.total_pages) if task.total_pages else "由系统决定"),
    ]
    for k, v in meta_lines:
        doc.add_paragraph(f"{k}：{v}", style="List Bullet")

    te = task.teaching_elements
    if isinstance(te, dict) and te:
        doc.add_paragraph()
        h = doc.add_paragraph()
        r = h.add_run("一、教学要素（结构化）")
        r.bold = True
        r.font.size = Pt(12)
        sections: list[tuple[str, Any]] = [
            ("核心知识点", te.get("knowledge_points")),
            ("教学目标", te.get("teaching_goals")),
            ("讲授逻辑 / 大纲", te.get("teaching_logic")),
            ("重点难点", te.get("key_difficulties")),
            ("课程时长", te.get("duration")),
            ("互动设计", te.get("interaction_design")),
            ("期望产出格式", te.get("output_formats")),
        ]
        for label, val in sections:
            for _, text in _te_lines(label, val):
                doc.add_paragraph(f"{label}：{text}")

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r = h2.add_run("二、教学说明与设计意图（任务描述摘要）")
    r.bold = True
    r.font.size = Pt(12)
    desc = (task.description or "").strip()
    if len(desc) > 12000:
        desc = desc[:11999] + "…"
    for chunk in desc.split("\n\n") if desc else ["（无）"]:
        doc.add_paragraph(chunk.strip() or "（空行）")

    doc.add_paragraph()
    h3 = doc.add_paragraph()
    r = h3.add_run("三、教学过程与课件分页要点")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(
        "以下由各页幻灯片内容抽取文本，供备课参考；具体呈现以课件为准。"
    )

    for i, pid in enumerate(task.page_order, start=1):
        page = task.pages.get(pid)
        if not page:
            continue
        inner = extract_slide_html_from_py(page.py_code or "")
        plain = _html_to_plain_text(inner, max_len=3500)
        doc.add_heading(f"第 {i} 页（{pid}）", level=2)
        if plain:
            doc.add_paragraph(plain)
        else:
            doc.add_paragraph("（本页无可用文本或为空）")

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "— 本文档由 EducationAgent PPT Agent 根据当前任务与课件自动生成 —"
    )
    for run in foot.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.save(str(output_path))
